from __future__ import annotations

from pathlib import Path
import re


class DocxDependencyError(RuntimeError):
    """Raised when python-docx is required but unavailable."""


def _load_document_module():
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise DocxDependencyError(
            "缺少 python-docx 依赖，无法读取或写入 docx。请执行 pip install -r requirements.txt。"
        ) from exc
    return Document, Pt, WD_ALIGN_PARAGRAPH, qn


def write_sections_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    Document, Pt, WD_ALIGN_PARAGRAPH, qn = _load_document_module()
    doc = Document()
    _configure_document_styles(doc, Pt, qn)
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for heading, paragraphs in sections:
        heading_paragraph = doc.add_heading(heading, level=1)
        heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for paragraph in paragraphs:
            _write_paragraph(doc, paragraph)
    doc.save(path)


def write_article_docx(path: Path, title: str, blocks: list[dict[str, str]]) -> None:
    Document, Pt, WD_ALIGN_PARAGRAPH, qn = _load_document_module()
    doc = Document()
    _configure_document_styles(doc, Pt, qn)
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for block in blocks:
        kind = block.get("type", "paragraph")
        text = block.get("text", "").strip()
        if not text:
            continue
        if kind == "heading":
            heading_paragraph = doc.add_heading(_strip_markdown(text), level=1)
            heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif kind == "subheading":
            subheading_paragraph = doc.add_heading(_strip_markdown(text), level=2)
            subheading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif kind == "list":
            _add_list_item(doc, text)
        else:
            paragraph = doc.add_paragraph()
            _write_inline_text(paragraph, text)
    doc.save(path)


def read_docx_text(path: Path) -> str:
    Document, _, _, _ = _load_document_module()
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def _configure_document_styles(doc, Pt, qn) -> None:
    normal_style = doc.styles["Normal"]
    _apply_style_font(normal_style, "Microsoft YaHei", Pt(11), qn)
    if hasattr(normal_style, "paragraph_format"):
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(8)
        normal_style.paragraph_format.first_line_indent = Pt(22)

    title_style = doc.styles["Title"]
    _apply_style_font(title_style, "Microsoft YaHei", Pt(18), qn, bold=True)
    if hasattr(title_style, "paragraph_format"):
        title_style.paragraph_format.space_after = Pt(14)
        title_style.paragraph_format.space_before = Pt(6)

    heading1_style = doc.styles["Heading 1"]
    _apply_style_font(heading1_style, "Microsoft YaHei", Pt(15), qn, bold=True)
    if hasattr(heading1_style, "paragraph_format"):
        heading1_style.paragraph_format.space_before = Pt(14)
        heading1_style.paragraph_format.space_after = Pt(8)

    heading2_style = doc.styles["Heading 2"]
    _apply_style_font(heading2_style, "Microsoft YaHei", Pt(13), qn, bold=True)
    if hasattr(heading2_style, "paragraph_format"):
        heading2_style.paragraph_format.space_before = Pt(10)
        heading2_style.paragraph_format.space_after = Pt(6)

    for style_name in ("List Bullet", "List Number"):
        list_style = doc.styles[style_name]
        _apply_style_font(list_style, "Microsoft YaHei", Pt(11), qn)
        if hasattr(list_style, "paragraph_format"):
            list_style.paragraph_format.space_after = Pt(6)


def _apply_style_font(style, font_name: str, font_size, qn, *, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = font_size
    style.font.bold = bold
    if style.element.rPr is not None:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _write_paragraph(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    _write_inline_text(paragraph, text)


def _add_list_item(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number" if re.match(r"^\d+[\.\、\)]\s*", text) else "List Bullet")
    _write_inline_text(paragraph, _strip_markdown(text))


def _write_inline_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(_strip_markdown(part))
            run.bold = False


def _strip_markdown(text: str) -> str:
    text = re.sub(r"^\s*#+\s*", "", text)
    text = re.sub(r"^\s*[-*]\s*", "", text)
    return text
